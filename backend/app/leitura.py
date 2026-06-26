"""
leitura.py — Leitura de arquivos (versão API).

Mesma lógica do leitura.py original — extração de texto de PDF, Word,
Excel e txt, byte a byte idêntica. Única adaptação: o original esperava
um arquivo do Streamlit (que expõe `.name`); aqui recebemos um arquivo
do FastAPI (`UploadFile`, que expõe `.filename`) e um stream de bytes
(`.file`, compatível com pdfplumber.open / Document / pd.ExcelFile, que
todos aceitam um objeto tipo arquivo).

Nenhuma regra de negócio mudou — isso é só troca de "casca" de objeto.
"""

import io

import pandas as pd
import pdfplumber
from docx import Document


def read_pdf(file_stream) -> str:
    """Extrai texto de PDF (inclui tabelas, comuns em propostas)."""
    try:
        text_parts = []
        with pdfplumber.open(file_stream) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                for table in page.extract_tables():
                    for row in table:
                        if row:
                            text_parts.append(" | ".join([str(c) if c else "" for c in row]))
        return "\n".join(text_parts)
    except Exception as e:
        return f"[ERRO ao ler PDF: {e}]"


def read_docx(file_stream) -> str:
    """Extrai texto de DOCX (Word), incluindo tabelas."""
    try:
        doc = Document(file_stream)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[ERRO ao ler DOCX: {e}]"


def read_xlsx(file_stream) -> str:
    """Extrai conteúdo de Excel como texto estruturado (aba por aba)."""
    try:
        xl = pd.ExcelFile(file_stream)
        parts = []
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(file_stream, sheet_name=sheet_name, header=None)
            df = df.dropna(how="all").dropna(axis=1, how="all")
            if df.empty:
                continue
            parts.append(f"\n=== Aba: {sheet_name} ===")
            for _, row in df.iterrows():
                row_values = [str(v) for v in row.values if pd.notna(v) and str(v).strip()]
                if row_values:
                    parts.append(" | ".join(row_values))
        return "\n".join(parts)
    except Exception as e:
        return f"[ERRO ao ler Excel: {e}]"


def read_uploaded_file(filename: str, raw_bytes: bytes) -> str:
    """
    Detecta o tipo do arquivo pelo nome e extrai o texto.

    Recebe o nome do arquivo e os bytes crus (já lidos do UploadFile do
    FastAPI pela camada de rotas) — assim esta função não depende de
    nenhum framework web, só de bytes puros, o que facilita testar
    isoladamente.
    """
    nome = filename.lower()
    stream = io.BytesIO(raw_bytes)

    if nome.endswith(".pdf"):
        return read_pdf(stream)
    elif nome.endswith(".docx"):
        return read_docx(stream)
    elif nome.endswith((".xlsx", ".xls", ".xlsm")):
        return read_xlsx(stream)
    elif nome.endswith(".txt"):
        return raw_bytes.decode("utf-8", errors="ignore")
    else:
        return f"[Formato não suportado: {nome}]"
