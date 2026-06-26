"""
etapa6.py — Etapa 6: Equalização Comercial.

A mais pesada e valiosa do modelo de 8 etapas (definição do v8).

O que faz:
1. CHECKPOINT antes de chamar a IA: pergunta taxa de desconto e moeda de
   referência via chatbox (mesmo padrão da anualização na Etapa 2). Sem
   essas duas respostas, a equalização não roda — ficam gravadas em
   estudo.taxa_desconto e estudo.regra_moeda.
2. Normaliza a base comum: usa o modelo de precificação (Etapa 1) para
   decidir como tornar os preços comparáveis entre fornecedores e contra
   o baseline.
3. Ramifica pelo modelo de precificação:
   - hora-homem          → equaliza por taxa horária x volume de referência
   - por_funcionario     → equaliza por custo/posto x headcount de referência
   - mensal_fixo         → comparação direta (mensal x12), sem ramificação extra
   - variavel_por_consumo→ exige volume de referência do baseline; se não
                           houver, marca premissa e segue com o que a
                           proposta declarou
   - tpq                 → equaliza item a item (valores unitários)
   - misto               → trata cada componente conforme seu próprio modelo
4. Isola on-tops (NÃO mistura no preço base):
   - on-top de escopo: vem do delta_escopo da Etapa 3 (itens adicionados/
     removidos/modificados no edital vs baseline)
   - on-top de desvio/inclusão/exclusão: vem das inclusões/exclusões de
     escopo e do flag nao_cumpre_mandatorio da Etapa 4 (proposta inclui algo
     que as outras não, ou não cumpre um mandatório — isso tem valor/custo
     que precisa ficar rastreável separadamente)
5. Aplica ajustes de frete/impostos/prazo/moeda quando a 4B sinalizar que
   não estão inclusos, usando a taxa_desconto e regra_moeda informadas.
6. Saída: preço equalizado por fornecedor + savings vs baseline. Grava em
   estudo.equalizacao_comercial. SEM saída formatada (Word/Excel) ainda —
   fica para etapa posterior.
"""

import json
import time
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from .ia import call_claude
from .erros import ErroEtapa

MAX_TOKENS_ETAPA6 = 8000


# ---------------------------------------------------------------------------
# Prompt
# ---------------------------------------------------------------------------

SYSTEM_EQUALIZACAO = """
Você é um especialista sênior em procurement produzindo a EQUALIZAÇÃO
COMERCIAL de um estudo de propostas. Você recebe dados já extraídos por
outras etapas (não releia documento bruto): preço atual (baseline), dados
comerciais de cada proposta, delta de escopo do edital vs baseline, e
inclusões/exclusões/não-conformidades de escopo por fornecedor (ótica
técnica). Recebe também a taxa de desconto e a moeda de referência
informadas pelo consultor.

Sua tarefa é tornar os preços COMPARÁVEIS entre si e contra o baseline,
isolando o que é "preço base equivalente" do que é "ajuste por escopo
diferente" — nunca misture os dois na mesma linha.

Responda SOMENTE com um objeto JSON válido, sem texto antes ou depois:

{
  "moeda_referencia": "<moeda usada na equalização — a informada pelo consultor>",
  "taxa_desconto_aplicada": <número, % informado pelo consultor>,
  "por_fornecedor": [
    {
      "fornecedor": "<nome>",
      "preco_base_equalizado": <número ou null — preço comparável, ANTES de on-tops>,
      "metodo_equalizacao": "<1-2 frases: como chegou nesse número, qual ramo do modelo de precificação usou>",
      "on_tops_escopo": [
        {"item": "<item do delta_escopo do edital>", "valor_estimado": <número ou null>, "direcao": "soma|subtrai", "origem": "edital_vs_baseline"}
      ],
      "on_tops_desvio": [
        {"item": "<inclusão/exclusão/não-conformidade desta proposta>", "valor_estimado": <número ou null>, "direcao": "soma|subtrai", "origem": "inclusao_escopo|exclusao_escopo|nao_conformidade_mandatoria"}
      ],
      "ajustes_condicoes": [
        {"tipo": "frete|impostos|prazo|moeda", "comentario": "<o que foi ajustado e por quê>", "valor_estimado": <número ou null>}
      ],
      "preco_total_equalizado": <número ou null — base + on-tops + ajustes>,
      "savings_vs_baseline": <número ou null — positivo = economia, negativo = aumento>,
      "savings_percentual": <número ou null>,
      "premissas": ["<premissa assumida para ESTE fornecedor>"],
      "faltantes": ["<dado que faltou para equalizar ESTE fornecedor>"]
    }
  ],
  "sintese_comparativa": "<parágrafo: como os fornecedores se posicionam em preço equalizado, sem recomendar qual escolher>",
  "premissas_gerais": ["<premissa que vale para a equalização como um todo>"],
  "faltantes_gerais": ["<limitação geral desta equalização>"]
}

Regras de ramificação por modelo de precificação (decida o método olhando
modelo_precificacao_ofertado de cada proposta, que pode diferir do detectado
na Etapa 1):
- hora-homem: equalize por taxa horária × volume de referência do baseline
  (se não houver volume de referência, registre em faltantes e equalize
  pela taxa unitária apenas, sinalizando a limitação).
- por_funcionario: equalize por custo/posto × headcount de referência do
  baseline; mesma ressalva se não houver headcount de referência.
- mensal_fixo: comparação direta, anualizando (mensal × 12) se necessário.
- variavel_por_consumo: precisa do volume de referência do baseline; sem
  ele, equalização fica marcada como aproximada e registrada em faltantes.
- tpq: equalize item a item pelos valores unitários informados.
- misto: trate cada componente do preço pelo método que lhe cabe, e explique
  em metodo_equalizacao como os componentes foram combinados.

Regras gerais:
- NUNCA misture on-top de escopo (origem: mudança edital vs baseline) com
  on-top de desvio (origem: diferença entre o que ESTA proposta oferece e o
  que as outras oferecem, ou não-conformidade mandatória). São rastreáveis
  separadamente mesmo que o efeito final no preço seja somado.
- Quando valor_estimado de um on-top não puder ser calculado com precisão,
  use null e explique a direção qualitativa (soma/subtrai) mesmo assim —
  não invente número preciso sem base.
- NÃO recomende qual fornecedor escolher — isso é Etapa 7. Aqui é só tornar
  os números comparáveis e mostrar o savings.
- Se uma proposta não tiver preco_total_proposto (da extração comercial),
  registre em faltantes que a equalização desse fornecedor ficou incompleta,
  mas ainda assim preencha o que for possível com os itens_precificados.
- LIMITE DE TAMANHO: on_tops_escopo, on_tops_desvio e ajustes_condicoes
  máximo 8 itens cada por fornecedor; premissas e faltantes (por fornecedor
  e gerais) máximo 5 itens cada.
  O JSON completo deve caber em 7000 tokens.
"""


# ---------------------------------------------------------------------------
# Montagem de contexto (lê do Estudo, não relê documentos brutos)
# ---------------------------------------------------------------------------

def _resumo_baseline_comercial_ctx(estudo) -> str:
    if not estudo.baseline:
        return "Baseline comercial: não disponível."
    com = estudo.baseline.get("comercial", {})
    return (
        f"Baseline comercial disponível.\n"
        f"Preço anual atual: {com.get('preco_anual_total')} "
        f"(base de anualização: {com.get('base_anualização', '—')})\n"
        f"Valores unitários de referência: "
        f"{json.dumps(com.get('valores_unitarios', []), ensure_ascii=False)}"
    )


def _resumo_delta_escopo_ctx(estudo) -> str:
    delta = (estudo.edital or {}).get("delta_escopo", {})
    if not delta.get("tem_baseline"):
        return "Delta de escopo (edital vs baseline): não disponível."
    linhas = ["DELTA DE ESCOPO (edital vs baseline) — fonte dos on-tops de escopo:"]
    for a in delta.get("adicionados", []):
        linhas.append(f"- ADICIONADO: {a.get('item','')} (impacto de custo: {a.get('impacto_custo','?')})")
    for r in delta.get("removidos", []):
        linhas.append(f"- REMOVIDO: {r.get('item','')} (impacto de custo: {r.get('impacto_custo','?')})")
    for m in delta.get("modificados", []):
        linhas.append(
            f"- MODIFICADO: {m.get('item','')} — {m.get('antes','?')} -> {m.get('depois','?')} "
            f"(impacto de custo: {m.get('impacto_custo','?')})"
        )
    return "\n".join(linhas)


def _resumo_propostas_tecnicas_ctx(estudo) -> str:
    propostas = estudo.propostas_tecnicas or []
    if not propostas:
        return "Propostas técnicas: nenhuma avaliação disponível."
    blocos = ["AVALIAÇÃO TÉCNICA POR FORNECEDOR — fonte dos on-tops de desvio:"]
    for p in propostas:
        linhas = [f"=== {p.get('fornecedor','?')} ==="]
        if p.get("nao_cumpre_mandatorio"):
            linhas.append(
                f"NÃO CUMPRE MANDATÓRIO: {', '.join(p.get('mandatorios_nao_cumpridos', []))}"
            )
        incl = p.get("inclusoes_escopo", [])
        if incl:
            linhas.append(f"Inclusões de escopo (só este fornecedor): {', '.join(incl)}")
        excl = p.get("exclusoes_escopo", [])
        if excl:
            linhas.append(f"Exclusões de escopo (este fornecedor deixa de fora): {', '.join(excl)}")
        blocos.append("\n".join(linhas))
    return "\n\n".join(blocos)


def _resumo_propostas_comerciais_ctx(estudo) -> str:
    propostas = estudo.propostas_comerciais or []
    if not propostas:
        return "Propostas comerciais: nenhuma extração disponível."
    blocos = ["DADOS COMERCIAIS POR FORNECEDOR:"]
    for p in propostas:
        linhas = [f"=== {p.get('fornecedor','?')} ==="]
        linhas.append(f"Modelo de precificação ofertado: {p.get('modelo_precificacao_ofertado','—')}")
        linhas.append(f"Moeda: {p.get('moeda','—')}")
        linhas.append(f"Preço total proposto: {p.get('preco_total_proposto')}")
        linhas.append(f"Base de anualização: {p.get('base_anualizacao','—')}")
        linhas.append(f"Impostos inclusos: {p.get('impostos_inclusos','—')} | Frete incluso: {p.get('frete_incluso','—')}")
        linhas.append(f"Condições de pagamento: {p.get('condicoes_pagamento','—')} | Reajuste: {p.get('reajuste','—')}")
        itens = p.get("itens_precificados", [])
        if itens:
            linhas.append(f"Itens precificados: {json.dumps(itens, ensure_ascii=False)}")
        blocos.append("\n".join(linhas))
    return "\n\n".join(blocos)


def _parse_resposta(resposta_bruta: str) -> dict:
    """Parser com fallback de truncamento (mesmo padrão das outras etapas)."""
    texto = resposta_bruta.strip()
    if texto.startswith("```"):
        linhas = texto.split("\n")
        texto = "\n".join(linhas[1:-1]).strip()
    try:
        return json.loads(texto)
    except json.JSONDecodeError:
        texto_cortado = texto.rstrip().rstrip(",")
        abertas = texto_cortado.count("{") - texto_cortado.count("}")
        colchetes = texto_cortado.count("[") - texto_cortado.count("]")
        fechamento = "]" * colchetes + "}" * abertas
        try:
            return json.loads(texto_cortado + fechamento)
        except json.JSONDecodeError:
            raise ValueError(
                "JSON inválido mesmo após tentativa de correção. "
                "Estudo com muitos fornecedores/itens — verifique o volume de dados."
            )


# ---------------------------------------------------------------------------
# Checkpoint: taxa de desconto e moeda (via chatbox, bloqueia o fluxo)
# ---------------------------------------------------------------------------

def precisa_checkpoint_etapa6(estudo) -> bool:
    """True se taxa_desconto ou regra_moeda ainda não foram informadas."""
    return estudo.taxa_desconto is None or estudo.regra_moeda is None


def confirmar_taxa_e_moeda(estudo, taxa_desconto: float, regra_moeda: str) -> str:
    """
    Chamada quando o consultor informa taxa de desconto e moeda via chatbox.
    Grava no Estudo e registra a premissa.
    """
    estudo.taxa_desconto = taxa_desconto
    estudo.regra_moeda = regra_moeda
    estudo.add_premissa(
        f"Equalização comercial usou taxa de desconto de {taxa_desconto}% "
        f"e moeda de referência {regra_moeda}, informadas pelo consultor."
    )
    return f"✅ Taxa de desconto **{taxa_desconto}%** e moeda **{regra_moeda}** registradas."


# ---------------------------------------------------------------------------
# Função principal
# ---------------------------------------------------------------------------

def rodar_etapa6(estudo) -> dict:
    """
    Executa a Etapa 6 completa. NÃO chama a IA se o checkpoint de taxa de
    desconto/moeda ainda não foi respondido — quem chama esta função deve
    checar precisa_checkpoint_etapa6(estudo) antes.

    Retorna dict com:
        - tem_dados : bool
        - analise   : dict completo (JSON do Claude) ou None
        - resumo    : texto legível pra UI
    """
    if precisa_checkpoint_etapa6(estudo):
        msg = "Taxa de desconto e/ou moeda de referência ainda não informadas — Etapa 6 não pode rodar."
        return {"tem_dados": False, "analise": None, "resumo": f"⚠️ {msg}"}

    if not estudo.propostas_comerciais:
        msg = "Nenhuma proposta comercial extraída (Etapa 4B) — Etapa 6 não pôde rodar."
        estudo.add_faltante(msg)
        estudo.equalizacao_comercial = None
        return {"tem_dados": False, "analise": None, "resumo": f"⚠️ {msg}"}

    n_forn = len(estudo.propostas_comerciais)
    if n_forn <= 4:
        _lim_ontops, _lim_pf = 5, 3
    elif n_forn <= 7:
        _lim_ontops, _lim_pf = 3, 2
    else:
        _lim_ontops, _lim_pf = 2, 1

    _instrucao_tamanho = (
        f"LIMITE DE TAMANHO OBRIGATÓRIO ({n_forn} fornecedores neste estudo): "
        f"on_tops_escopo, on_tops_desvio e ajustes_condicoes máximo {_lim_ontops} itens cada por fornecedor; "
        f"premissas e faltantes (por fornecedor e gerais) máximo {_lim_pf} cada. "
        f"O JSON COMPLETO deve caber em 7000 tokens — seja conciso nas descrições, "
        f"priorize os campos numéricos (preços, savings) sobre texto explicativo."
    )

    contexto = (
        f"{_instrucao_tamanho}\n\n"
        f"Categoria: {estudo.categoria}\n"
        f"Modelo de precificação detectado (Etapa 1): {estudo.modelo_precificacao}\n"
        f"Micro-categoria: {estudo.micro_categoria or '—'}\n"
        f"Taxa de desconto informada pelo consultor: {estudo.taxa_desconto}%\n"
        f"Moeda de referência informada pelo consultor: {estudo.regra_moeda}\n\n"
        f"{_resumo_baseline_comercial_ctx(estudo)}\n\n"
        f"{_resumo_delta_escopo_ctx(estudo)}\n\n"
        f"{_resumo_propostas_tecnicas_ctx(estudo)}\n\n"
        f"{_resumo_propostas_comerciais_ctx(estudo)}"
    )

    print(f"[PERF] etapa6 início — 1 chamada para {n_forn} fornecedor(es) (chamada única), limites: {_lim_ontops} on-tops / {_lim_pf} pf por fornecedor")
    _t0 = time.time()
    resposta_bruta = call_claude(
        messages=[{"role": "user", "content": contexto}],
        system=SYSTEM_EQUALIZACAO,
        max_tokens=MAX_TOKENS_ETAPA6,
    )
    print(f"[PERF] etapa6 call_claude ({n_forn} fornecedor(es)): {time.time() - _t0:.1f}s")

    try:
        analise = _parse_resposta(resposta_bruta)
    except (json.JSONDecodeError, ValueError) as e:
        raise ErroEtapa(f"Erro ao interpretar resposta da IA na Etapa 6: {e}", resposta_bruta=resposta_bruta)

    estudo.equalizacao_comercial = analise

    for p in analise.get("premissas_gerais", []):
        estudo.add_premissa(p)
    for f in analise.get("faltantes_gerais", []):
        estudo.add_faltante(f)
    for forn in analise.get("por_fornecedor", []):
        nome = forn.get("fornecedor", "?")
        for p in forn.get("premissas", []):
            estudo.add_premissa(f"[{nome}] {p}")
        for f in forn.get("faltantes", []):
            estudo.add_faltante(f"[{nome}] {f}")

    estudo.etapa_atual = 6

    resumo = _montar_resumo(analise)
    return {"tem_dados": True, "analise": analise, "resumo": resumo}


def _montar_resumo(analise: dict) -> str:
    linhas = []

    linhas.append(
        f"**Equalização comercial** — moeda: {analise.get('moeda_referencia','—')} · "
        f"taxa de desconto aplicada: {analise.get('taxa_desconto_aplicada','—')}%"
    )

    sintese = analise.get("sintese_comparativa")
    if sintese:
        linhas.append(f"\n**Síntese comparativa:** {sintese}")

    for forn in analise.get("por_fornecedor", []):
        linhas.append(f"\n---\n### {forn.get('fornecedor','?')}")

        base = forn.get("preco_base_equalizado")
        total = forn.get("preco_total_equalizado")
        savings = forn.get("savings_vs_baseline")
        savings_pct = forn.get("savings_percentual")

        if base is not None:
            linhas.append(f"**Preço base equalizado:** {base:,.2f}")
        linhas.append(f"_Método: {forn.get('metodo_equalizacao','—')}_")

        on_tops_escopo = forn.get("on_tops_escopo", [])
        if on_tops_escopo:
            linhas.append("\n**On-tops de escopo (edital vs baseline):**")
            for o in on_tops_escopo:
                sinal = "+" if o.get("direcao") == "soma" else "−"
                valor = o.get("valor_estimado")
                valor_str = f"{sinal}{valor:,.2f}" if valor is not None else f"{sinal}(não estimado)"
                linhas.append(f"- {o.get('item','')}: {valor_str}")

        on_tops_desvio = forn.get("on_tops_desvio", [])
        if on_tops_desvio:
            linhas.append("\n**On-tops de desvio (inclusão/exclusão/não-conformidade):**")
            for o in on_tops_desvio:
                sinal = "+" if o.get("direcao") == "soma" else "−"
                valor = o.get("valor_estimado")
                valor_str = f"{sinal}{valor:,.2f}" if valor is not None else f"{sinal}(não estimado)"
                linhas.append(f"- [{o.get('origem','?')}] {o.get('item','')}: {valor_str}")

        ajustes = forn.get("ajustes_condicoes", [])
        if ajustes:
            linhas.append("\n**Ajustes de condições (frete/impostos/prazo/moeda):**")
            for a in ajustes:
                valor = a.get("valor_estimado")
                valor_str = f"{valor:,.2f}" if valor is not None else "(não estimado)"
                linhas.append(f"- [{a.get('tipo','?')}] {a.get('comentario','')}: {valor_str}")

        if total is not None:
            linhas.append(f"\n**Preço total equalizado:** {total:,.2f}")
        if savings is not None:
            sinal_s = "economia" if savings >= 0 else "aumento"
            pct_str = f" ({savings_pct:.1f}%)" if savings_pct is not None else ""
            linhas.append(f"**Savings vs baseline:** {sinal_s} de {abs(savings):,.2f}{pct_str}")

    faltantes = analise.get("faltantes_gerais", [])
    if faltantes:
        linhas.append("\n**Limitações desta equalização:**")
        for f in faltantes:
            linhas.append(f"- {f}")

    return "\n".join(linhas)


# ---------------------------------------------------------------------------
# Geração Word (one-pager) — mesmo estilo visual das Etapas 5 e 7
# ---------------------------------------------------------------------------

COR_AZUL_AM = RGBColor(0x1F, 0x3A, 0x5F)
COR_CINZA = RGBColor(0x59, 0x59, 0x59)
COR_VERDE = RGBColor(0x1E, 0x7A, 0x3C)
COR_VERMELHO = RGBColor(0xC0, 0x39, 0x2B)


def _heading(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(texto)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COR_AZUL_AM


def _linha_divisoria(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _sombrear_celula(cell, cor_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor_hex)
    tcPr.append(shd)


def gerar_word_etapa6(estudo) -> BytesIO:
    """Gera o one-pager Word da Etapa 6. Retorna BytesIO pronto para download."""
    analise = estudo.equalizacao_comercial or {}
    moeda = analise.get("moeda_referencia", "—")
    taxa = analise.get("taxa_desconto_aplicada", "—")
    por_forn = analise.get("por_fornecedor", [])
    n_forn = len(por_forn)

    doc = Document()
    secao = doc.sections[0]
    secao.left_margin = Cm(1.8)
    secao.right_margin = Cm(1.8)
    secao.top_margin = Cm(1.5)
    secao.bottom_margin = Cm(1.5)

    estilo = doc.styles["Normal"]
    estilo.font.name = "Arial"
    estilo.font.size = Pt(10)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = titulo.add_run("Equalização Comercial — One-Pager")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COR_AZUL_AM

    sub = doc.add_paragraph()
    run_sub = sub.add_run(
        f"{estudo.micro_categoria or estudo.categoria or 'Categoria não identificada'} · "
        f"{n_forn} fornecedor(es) · moeda {moeda} · taxa de desconto {taxa}%"
    )
    run_sub.font.size = Pt(10)
    run_sub.font.color.rgb = COR_CINZA
    run_sub.italic = True

    _linha_divisoria(doc)

    # Síntese comparativa
    sintese = analise.get("sintese_comparativa")
    if sintese:
        _heading(doc, "Síntese Comparativa")
        doc.add_paragraph(sintese)

    # Tabela de savings por fornecedor
    if por_forn:
        _heading(doc, "Síntese de Savings por Fornecedor")
        n_cols = 5
        tabela = doc.add_table(rows=1, cols=n_cols)
        tabela.style = "Table Grid"
        tabela.alignment = WD_TABLE_ALIGNMENT.LEFT
        cabecalho = [
            "Fornecedor",
            f"Preço Base Equalizado ({moeda})",
            f"Preço Total Equalizado ({moeda})",
            f"Savings ({moeda})",
            "Savings %",
        ]
        hdr_cells = tabela.rows[0].cells
        for i, txt in enumerate(cabecalho):
            hdr_cells[i].text = ""
            run = hdr_cells[i].paragraphs[0].add_run(txt)
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _sombrear_celula(hdr_cells[i], "1F3A5F")

        for forn in por_forn:
            row = tabela.add_row().cells
            row[0].text = ""
            run_nome = row[0].paragraphs[0].add_run(forn.get("fornecedor", "?"))
            run_nome.font.size = Pt(8)
            run_nome.font.bold = True

            base = forn.get("preco_base_equalizado")
            total = forn.get("preco_total_equalizado")
            savings = forn.get("savings_vs_baseline")
            savings_pct = forn.get("savings_percentual")

            row[1].text = ""
            row[1].paragraphs[0].add_run(f"{base:,.2f}" if base is not None else "—").font.size = Pt(8)
            row[2].text = ""
            row[2].paragraphs[0].add_run(f"{total:,.2f}" if total is not None else "—").font.size = Pt(8)

            row[3].text = ""
            run_sav = row[3].paragraphs[0].add_run(f"{savings:,.2f}" if savings is not None else "—")
            run_sav.font.size = Pt(8)
            if savings is not None:
                run_sav.font.color.rgb = COR_VERDE if savings >= 0 else COR_VERMELHO

            row[4].text = ""
            run_pct = row[4].paragraphs[0].add_run(f"{savings_pct:.1f}%" if savings_pct is not None else "—")
            run_pct.font.size = Pt(8)
            if savings_pct is not None:
                run_pct.font.color.rgb = COR_VERDE if savings_pct >= 0 else COR_VERMELHO

        larguras = [Cm(4.5), Cm(3.5), Cm(3.5), Cm(3.0), Cm(2.0)]
        for row in tabela.rows:
            for i, cell in enumerate(row.cells):
                cell.width = larguras[i]

    # On-tops e ajustes por fornecedor
    tem_ontops = any(
        forn.get("on_tops_escopo") or forn.get("on_tops_desvio") or forn.get("ajustes_condicoes")
        for forn in por_forn
    )
    if tem_ontops:
        _heading(doc, "On-tops e Ajustes por Fornecedor")
        for forn in por_forn:
            on_tops_escopo = forn.get("on_tops_escopo", [])
            on_tops_desvio = forn.get("on_tops_desvio", [])
            ajustes = forn.get("ajustes_condicoes", [])
            if not on_tops_escopo and not on_tops_desvio and not ajustes:
                continue
            p = doc.add_paragraph()
            p.add_run(forn.get("fornecedor", "?")).bold = True
            p.paragraph_format.space_before = Pt(6)
            for o in on_tops_escopo:
                sinal = "+" if o.get("direcao") == "soma" else "−"
                valor = o.get("valor_estimado")
                valor_str = f"{sinal}{valor:,.2f}" if valor is not None else f"{sinal}(não estimado)"
                doc.add_paragraph(f"[escopo] {o.get('item','')} — {valor_str}", style="List Bullet")
            for o in on_tops_desvio:
                sinal = "+" if o.get("direcao") == "soma" else "−"
                valor = o.get("valor_estimado")
                valor_str = f"{sinal}{valor:,.2f}" if valor is not None else f"{sinal}(não estimado)"
                doc.add_paragraph(
                    f"[{o.get('origem','desvio')}] {o.get('item','')} — {valor_str}",
                    style="List Bullet",
                )
            for a in ajustes:
                valor = a.get("valor_estimado")
                valor_str = f"{valor:,.2f}" if valor is not None else "(não estimado)"
                doc.add_paragraph(
                    f"[{a.get('tipo','?')}] {a.get('comentario','')} — {valor_str}",
                    style="List Bullet",
                )

    # Premissas e limitações
    premissas_gerais = analise.get("premissas_gerais", [])
    faltantes_gerais = analise.get("faltantes_gerais", [])
    if premissas_gerais or faltantes_gerais:
        _heading(doc, "Premissas e Limitações desta Equalização")
        if premissas_gerais:
            doc.add_paragraph().add_run("Premissas:").bold = True
            for pr in premissas_gerais:
                doc.add_paragraph(pr, style="List Bullet")
        if faltantes_gerais:
            doc.add_paragraph().add_run("Limitações:").bold = True
            for f in faltantes_gerais:
                doc.add_paragraph(f, style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Geração Excel (3 abas) — mesmo estilo visual das Etapas 5 e 7
# ---------------------------------------------------------------------------

FILL_HEADER_E6 = PatternFill(start_color="1F3A5F", end_color="1F3A5F", fill_type="solid")
FONT_HEADER_E6 = Font(name="Arial", size=10, bold=True, color="FFFFFF")
FONT_BODY_E6 = Font(name="Arial", size=10)
FONT_BODY_BOLD_E6 = Font(name="Arial", size=10, bold=True)
THIN_E6 = Side(style="thin", color="CCCCCC")
BORDA_E6 = Border(top=THIN_E6, bottom=THIN_E6, left=THIN_E6, right=THIN_E6)
FILL_VERDE_E6 = PatternFill(start_color="C6E8C6", end_color="C6E8C6", fill_type="solid")
FILL_VERMELHO_E6 = PatternFill(start_color="F4C7C3", end_color="F4C7C3", fill_type="solid")


def _cabecalho_aba_e6(ws, colunas):
    for i, texto in enumerate(colunas, start=1):
        cell = ws.cell(row=1, column=i, value=texto)
        cell.font = FONT_HEADER_E6
        cell.fill = FILL_HEADER_E6
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = BORDA_E6
    ws.row_dimensions[1].height = 28
    ws.freeze_panes = "A2"


def gerar_excel_etapa6(estudo) -> BytesIO:
    """Gera o Excel da Etapa 6 (3 abas). Retorna BytesIO pronto para download."""
    analise = estudo.equalizacao_comercial or {}
    moeda = analise.get("moeda_referencia", "—")

    wb = Workbook()
    _aba_sintese_savings(wb, analise, moeda)
    _aba_ontops_ajustes(wb, analise, moeda)
    _aba_premissas_limitacoes(wb, analise)

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def _aba_sintese_savings(wb, analise, moeda):
    ws = wb.active
    ws.title = "Síntese de Savings"
    colunas = [
        "Fornecedor",
        "Método de Equalização",
        f"Preço Base Equalizado ({moeda})",
        f"Preço Total Equalizado ({moeda})",
        f"Savings vs Baseline ({moeda})",
        "Savings %",
    ]
    _cabecalho_aba_e6(ws, colunas)

    for r, forn in enumerate(analise.get("por_fornecedor", []), start=2):
        ws.cell(row=r, column=1, value=forn.get("fornecedor", "?")).font = FONT_BODY_BOLD_E6
        c_met = ws.cell(row=r, column=2, value=forn.get("metodo_equalizacao", "—"))
        c_met.font = FONT_BODY_E6
        c_met.alignment = Alignment(wrap_text=True, vertical="top")

        base = forn.get("preco_base_equalizado")
        total = forn.get("preco_total_equalizado")
        savings = forn.get("savings_vs_baseline")
        savings_pct = forn.get("savings_percentual")

        c_base = ws.cell(row=r, column=3, value=base)
        c_base.font = FONT_BODY_E6
        c_base.number_format = "#,##0.00"

        c_total = ws.cell(row=r, column=4, value=total)
        c_total.font = FONT_BODY_E6
        c_total.number_format = "#,##0.00"

        c_sav = ws.cell(row=r, column=5, value=savings)
        c_sav.font = FONT_BODY_E6
        c_sav.number_format = "#,##0.00"
        if savings is not None:
            c_sav.fill = FILL_VERDE_E6 if savings >= 0 else FILL_VERMELHO_E6

        c_pct = ws.cell(row=r, column=6, value=savings_pct)
        c_pct.font = FONT_BODY_E6
        c_pct.number_format = '0.0"%"'
        if savings_pct is not None:
            c_pct.fill = FILL_VERDE_E6 if savings_pct >= 0 else FILL_VERMELHO_E6

        for c in range(1, 7):
            ws.cell(row=r, column=c).border = BORDA_E6
            if c not in (2,):
                ws.cell(row=r, column=c).alignment = Alignment(horizontal="center", vertical="center")

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 38
    ws.column_dimensions["C"].width = 24
    ws.column_dimensions["D"].width = 24
    ws.column_dimensions["E"].width = 24
    ws.column_dimensions["F"].width = 12


def _aba_ontops_ajustes(wb, analise, moeda):
    ws = wb.create_sheet("On-tops e Ajustes")
    colunas = ["Fornecedor", "Categoria", "Item / Comentário", "Direção", f"Valor Estimado ({moeda})", "Origem"]
    _cabecalho_aba_e6(ws, colunas)

    r = 2
    for forn in analise.get("por_fornecedor", []):
        nome = forn.get("fornecedor", "?")
        for o in forn.get("on_tops_escopo", []):
            ws.cell(row=r, column=1, value=nome).font = FONT_BODY_BOLD_E6
            ws.cell(row=r, column=2, value="on-top de escopo").font = FONT_BODY_E6
            ws.cell(row=r, column=3, value=o.get("item", "")).font = FONT_BODY_E6
            ws.cell(row=r, column=4, value=o.get("direcao", "—")).font = FONT_BODY_E6
            c_val = ws.cell(row=r, column=5, value=o.get("valor_estimado"))
            c_val.font = FONT_BODY_E6
            c_val.number_format = "#,##0.00"
            ws.cell(row=r, column=6, value=o.get("origem", "edital_vs_baseline")).font = FONT_BODY_E6
            for c in range(1, 7):
                ws.cell(row=r, column=c).border = BORDA_E6
                ws.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="top")
            r += 1
        for o in forn.get("on_tops_desvio", []):
            ws.cell(row=r, column=1, value=nome).font = FONT_BODY_BOLD_E6
            ws.cell(row=r, column=2, value="on-top de desvio").font = FONT_BODY_E6
            ws.cell(row=r, column=3, value=o.get("item", "")).font = FONT_BODY_E6
            ws.cell(row=r, column=4, value=o.get("direcao", "—")).font = FONT_BODY_E6
            c_val = ws.cell(row=r, column=5, value=o.get("valor_estimado"))
            c_val.font = FONT_BODY_E6
            c_val.number_format = "#,##0.00"
            ws.cell(row=r, column=6, value=o.get("origem", "—")).font = FONT_BODY_E6
            for c in range(1, 7):
                ws.cell(row=r, column=c).border = BORDA_E6
                ws.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="top")
            r += 1
        for a in forn.get("ajustes_condicoes", []):
            ws.cell(row=r, column=1, value=nome).font = FONT_BODY_BOLD_E6
            ws.cell(row=r, column=2, value="ajuste de condição").font = FONT_BODY_E6
            ws.cell(row=r, column=3, value=a.get("comentario", "")).font = FONT_BODY_E6
            ws.cell(row=r, column=4, value="—").font = FONT_BODY_E6
            c_val = ws.cell(row=r, column=5, value=a.get("valor_estimado"))
            c_val.font = FONT_BODY_E6
            c_val.number_format = "#,##0.00"
            ws.cell(row=r, column=6, value=a.get("tipo", "—")).font = FONT_BODY_E6
            for c in range(1, 7):
                ws.cell(row=r, column=c).border = BORDA_E6
                ws.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="top")
            r += 1

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 20
    ws.column_dimensions["C"].width = 42
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 24
    ws.column_dimensions["F"].width = 28


def _aba_premissas_limitacoes(wb, analise):
    ws = wb.create_sheet("Premissas e Limitações")
    colunas = ["Tipo", "Fornecedor", "Texto"]
    _cabecalho_aba_e6(ws, colunas)

    r = 2
    for pr in analise.get("premissas_gerais", []):
        ws.cell(row=r, column=1, value="Premissa geral").font = FONT_BODY_E6
        ws.cell(row=r, column=2, value="—").font = FONT_BODY_E6
        ws.cell(row=r, column=3, value=pr).font = FONT_BODY_E6
        for c in range(1, 4):
            ws.cell(row=r, column=c).border = BORDA_E6
            ws.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="top")
        r += 1
    for f in analise.get("faltantes_gerais", []):
        ws.cell(row=r, column=1, value="Limitação geral").font = FONT_BODY_E6
        ws.cell(row=r, column=2, value="—").font = FONT_BODY_E6
        ws.cell(row=r, column=3, value=f).font = FONT_BODY_E6
        for c in range(1, 4):
            ws.cell(row=r, column=c).border = BORDA_E6
            ws.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="top")
        r += 1
    for forn in analise.get("por_fornecedor", []):
        nome = forn.get("fornecedor", "?")
        for pr in forn.get("premissas", []):
            ws.cell(row=r, column=1, value="Premissa por fornecedor").font = FONT_BODY_E6
            ws.cell(row=r, column=2, value=nome).font = FONT_BODY_BOLD_E6
            ws.cell(row=r, column=3, value=pr).font = FONT_BODY_E6
            for c in range(1, 4):
                ws.cell(row=r, column=c).border = BORDA_E6
                ws.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="top")
            r += 1
        for f in forn.get("faltantes", []):
            ws.cell(row=r, column=1, value="Limitação por fornecedor").font = FONT_BODY_E6
            ws.cell(row=r, column=2, value=nome).font = FONT_BODY_BOLD_E6
            ws.cell(row=r, column=3, value=f).font = FONT_BODY_E6
            for c in range(1, 4):
                ws.cell(row=r, column=c).border = BORDA_E6
                ws.cell(row=r, column=c).alignment = Alignment(wrap_text=True, vertical="top")
            r += 1

    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["C"].width = 60
